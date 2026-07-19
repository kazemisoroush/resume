#!/usr/bin/env python3
"""Render the résumé site and PDF from content.yaml.

content.yaml is the single source of truth. This script generates:
  - index.html         the portfolio homepage (with schema.org Person data)
  - resume.json        the résumé in the JSON Resume schema
  - resume.txt         a plain-text résumé
  - llms.txt           a discovery file for AI agents
  - resume.docx        a Word résumé (if python-docx is installed)
  - resume-print.html  a print-format document
  - resume.pdf         a tagged PDF from resume-print.html via WeasyPrint (if installed)

Contact details are read from the environment (PHONE_NUMBER / EMAIL_ADDRESS /
HOMEPAGE) at build time, so nothing personal lives in the source. Any value that
is unset is simply omitted.

Usage: python build.py
"""

import html
import json
import os
from pathlib import Path

import yaml

HERE = Path(__file__).parent


def h(s):
    return html.escape(str(s), quote=False)


def read_contact():
    return {
        "email": os.environ.get("EMAIL_ADDRESS", "").strip(),
        "homepage": os.environ.get("HOMEPAGE", "").strip(),
        "phone": os.environ.get("PHONE_NUMBER", "").strip(),
    }


def strip_scheme(url):
    return str(url).replace("https://", "").replace("http://", "").rstrip("/")


def homepage_url(contact):
    return "https://" + strip_scheme(contact["homepage"])


def is_present(end):
    return str(end).strip().lower() == "present"


COUNTRY_CODE = {"Australia": "AU", "Iran": "IR"}


def ccode(country):
    return COUNTRY_CODE.get(country, str(country)[:2].upper())


def split_commas(s):
    """Split on top-level commas, keeping commas inside parentheses together."""
    out, depth, cur = [], 0, ""
    for ch in str(s):
        if ch == "(":
            depth += 1
            cur += ch
        elif ch == ")":
            depth -= 1
            cur += ch
        elif ch == "," and depth == 0:
            out.append(cur.strip())
            cur = ""
        else:
            cur += ch
    if cur.strip():
        out.append(cur.strip())
    return out


def stack(tech):
    return " · ".join(split_commas(tech))


MONTHS = {"jan": "01", "feb": "02", "mar": "03", "apr": "04", "may": "05",
          "jun": "06", "jul": "07", "aug": "08", "sep": "09", "oct": "10",
          "nov": "11", "dec": "12"}


def iso_date(s):
    """'Jan 2024' -> '2024-01', '2014' -> '2014', 'Present' -> ''."""
    s = str(s).strip()
    if is_present(s):
        return ""
    parts = s.split()
    if len(parts) == 2 and parts[0][:3].lower() in MONTHS:
        return parts[1] + "-" + MONTHS[parts[0][:3].lower()]
    return s


def split_location(loc):
    """'Sydney, NSW' -> ('Sydney', 'NSW')."""
    bits = [b.strip() for b in str(loc).split(",")]
    return bits[0], (bits[1] if len(bits) > 1 else "")


def skill_keywords(data):
    kws = []
    for sk in data["skills"]:
        kws.extend(split_commas(sk["items"]))
    return kws


BASE_URL = "https://kazemisoroush.github.io/resume/"


# ===========================================================================
# Structured data (for AI, agents, scrapers, and ATS)
# ===========================================================================

def render_json_resume(data, contact):
    """The résumé in the JSON Resume schema (https://jsonresume.org)."""
    city, region = split_location(data["location"])
    basics = {
        "name": data["name"],
        "label": data["title"],
        "summary": " ".join(data["about"]),
        "url": BASE_URL,
        "location": {"city": city, "region": region, "countryCode": ccode(data["country"])},
        "profiles": [],
    }
    if contact["email"]:
        basics["email"] = contact["email"]
    if contact["phone"]:
        basics["phone"] = contact["phone"]
    if contact["homepage"]:
        hp = strip_scheme(contact["homepage"])
        basics["profiles"].append({"network": "LinkedIn", "url": homepage_url(contact), "username": hp.split("/")[-1]})
    basics["profiles"].append({"network": "GitHub", "url": data["github"], "username": strip_scheme(data["github"]).split("/")[-1]})

    work = []
    for e in data["experience"]:
        item = {
            "name": e["org"],
            "position": e["role"],
            "location": e["city"] + ", " + e["country"],
            "startDate": iso_date(e["start"]),
            "summary": e["summary"],
            "highlights": list(e.get("bullets") or []),
        }
        end = iso_date(e["end"])
        if end:
            item["endDate"] = end
        if e.get("tech"):
            item["keywords"] = split_commas(e["tech"])
        work.append(item)

    education = [{
        "institution": ed["org"],
        "studyType": ed["degree"],
        "area": ed["detail"],
        "startDate": iso_date(ed["start"]),
        "endDate": iso_date(ed["end"]),
        "location": ed["city"] + ", " + ed["country"],
    } for ed in data["education"]]

    skills = [{"name": sk["label"], "keywords": split_commas(sk["items"])} for sk in data["skills"]]

    resume = {
        "$schema": "https://raw.githubusercontent.com/jsonresume/resume-schema/v1.0.0/schema.json",
        "basics": basics,
        "work": work,
        "education": education,
        "skills": skills,
        "meta": {"canonical": BASE_URL + "resume.json", "version": "1.0.0"},
    }
    return json.dumps(resume, indent=2, ensure_ascii=False) + "\n"


def render_jsonld(data, contact):
    """A schema.org Person block for the homepage."""
    city, region = split_location(data["location"])
    person = {
        "@context": "https://schema.org",
        "@type": "Person",
        "name": data["name"],
        "jobTitle": data["title"],
        "worksFor": {"@type": "Organization", "name": data["experience"][0]["org"]},
        "address": {"@type": "PostalAddress", "addressLocality": city, "addressRegion": region, "addressCountry": ccode(data["country"])},
        "url": BASE_URL,
        "alumniOf": {"@type": "CollegeOrUniversity", "name": data["education"][0]["org"]},
        "knowsAbout": skill_keywords(data),
        "sameAs": [],
    }
    if contact["email"]:
        person["email"] = contact["email"]
    if contact["homepage"]:
        person["sameAs"].append(homepage_url(contact))
    person["sameAs"].append(data["github"])
    payload = json.dumps(person, indent=2, ensure_ascii=False)
    payload = payload.replace("<", "\\u003c").replace(">", "\\u003e").replace("&", "\\u0026")
    return '<script type="application/ld+json">\n' + payload + '\n</script>\n'


def render_text(data, contact):
    """A plain-text résumé, laid out top to bottom for any parser."""
    out = [data["name"], data["title"], data["location"] + ", " + data["country"]]
    row = []
    if contact["phone"]:
        row.append(contact["phone"])
    if contact["email"]:
        row.append(contact["email"])
    if contact["homepage"]:
        row.append(strip_scheme(contact["homepage"]))
    row.append(strip_scheme(data["github"]))
    out.append(" | ".join(row))
    out += ["", "SUMMARY"] + list(data["about"])
    out += ["", "EXPERIENCE"]
    for e in data["experience"]:
        end = "Present" if is_present(e["end"]) else e["end"]
        out.append(e["role"] + ", " + e["org"])
        out.append(str(e["start"]) + " - " + str(end) + " | " + e["city"] + ", " + e["country"])
        out.append(e["summary"])
        for b in (e.get("bullets") or []):
            out.append("- " + b)
        if e.get("tech"):
            out.append("Tech: " + ", ".join(split_commas(e["tech"])))
        out.append("")
    out.append("SKILLS")
    for sk in data["skills"]:
        out.append(sk["label"] + ": " + sk["items"])
    out += ["", "EDUCATION"]
    for ed in data["education"]:
        out.append(ed["degree"] + ", " + ed["org"])
        out.append(str(ed["start"]) + " - " + str(ed["end"]) + " | " + ed["detail"])
    return "\n".join(out).rstrip() + "\n"


def render_llms(data, contact):
    """An llms.txt so AI agents can find the structured résumé."""
    exp0 = data["experience"][0]
    out = ["# " + data["name"] + ", " + data["title"], "",
           exp0["role"] + " at " + exp0["org"] + ", based in " + data["location"] + ", " + data["country"] + ".",
           "", "## Résumé files",
           "- [PDF](" + BASE_URL + "resume.pdf)",
           "- [JSON Resume](" + BASE_URL + "resume.json)",
           "- [Plain text](" + BASE_URL + "resume.txt)",
           "- [Homepage](" + BASE_URL + ")",
           "", "## Summary"] + list(data["about"])
    return "\n".join(out) + "\n"


# ===========================================================================
# Homepage (GitHub Pages)
# ===========================================================================

SITE_CSS = """
  :root {
    --bg: #eef0f3; --surface: #ffffff; --surface-2: #f6f7f9;
    --ink: #14171d; --ink-2: #3b414d; --muted: #6b7280; --faint: #9aa1ac;
    --line: #dde1e7; --line-strong: #cbd0d9;
    --accent: #b06b1e; --accent-ink: #ffffff; --accent-soft: rgba(176,107,30,.12);
    --shadow: 0 1px 2px rgba(20,23,29,.04), 0 10px 30px rgba(20,23,29,.06);
    --font-display: ui-serif, "New York", "Iowan Old Style", Palatino, Georgia, serif;
    --font-sans: -apple-system, system-ui, "Segoe UI", Roboto, Helvetica, Arial, sans-serif;
    --font-mono: ui-monospace, "SF Mono", "JetBrains Mono", Menlo, Consolas, monospace;
    --measure: 800px;
  }
  @media (prefers-color-scheme: dark) {
    :root {
      --bg: #0d0f13; --surface: #15181e; --surface-2: #1a1e25;
      --ink: #e9ebee; --ink-2: #c2c7d0; --muted: #8b93a1; --faint: #667080;
      --line: #262b33; --line-strong: #333a44;
      --accent: #e0a15c; --accent-ink: #14171d; --accent-soft: rgba(224,161,92,.14);
      --shadow: 0 1px 2px rgba(0,0,0,.4), 0 12px 34px rgba(0,0,0,.5);
    }
  }
  :root[data-theme="light"] {
    --bg: #eef0f3; --surface: #ffffff; --surface-2: #f6f7f9;
    --ink: #14171d; --ink-2: #3b414d; --muted: #6b7280; --faint: #9aa1ac;
    --line: #dde1e7; --line-strong: #cbd0d9;
    --accent: #b06b1e; --accent-ink: #ffffff; --accent-soft: rgba(176,107,30,.12);
    --shadow: 0 1px 2px rgba(20,23,29,.04), 0 10px 30px rgba(20,23,29,.06);
  }
  :root[data-theme="dark"] {
    --bg: #0d0f13; --surface: #15181e; --surface-2: #1a1e25;
    --ink: #e9ebee; --ink-2: #c2c7d0; --muted: #8b93a1; --faint: #667080;
    --line: #262b33; --line-strong: #333a44;
    --accent: #e0a15c; --accent-ink: #14171d; --accent-soft: rgba(224,161,92,.14);
    --shadow: 0 1px 2px rgba(0,0,0,.4), 0 12px 34px rgba(0,0,0,.5);
  }
  * { box-sizing: border-box; }
  html { scroll-behavior: smooth; }
  body {
    margin: 0; background: var(--bg); color: var(--ink);
    font-family: var(--font-sans); font-size: 16.5px; line-height: 1.62;
    -webkit-font-smoothing: antialiased; text-rendering: optimizeLegibility;
  }
  .wrap { max-width: var(--measure); margin: 0 auto; padding: 0 28px; }
  a { color: inherit; text-decoration: none; }
  .label { display: inline-flex; align-items: center; gap: 10px; font-family: var(--font-mono); font-size: 11.5px; letter-spacing: .18em; text-transform: uppercase; color: var(--muted); font-weight: 500; }
  .label::before { content: ""; width: 22px; height: 2px; background: var(--accent); border-radius: 2px; flex: none; }
  .rule-head { display: grid; grid-template-columns: auto 1fr; align-items: center; gap: 18px; margin: 0 0 34px; }
  .rule-head::after { content: ""; height: 1px; background: var(--line); }
  header.hero { padding: 84px 0 40px; }
  .hero-grid { display: grid; grid-template-columns: 1fr auto; gap: 44px; align-items: center; }
  .portrait { width: 172px; height: 215px; object-fit: cover; object-position: 50% 22%; border-radius: 16px; border: 1px solid var(--line); box-shadow: var(--shadow); flex: none; }
  .kicker { font-family: var(--font-mono); font-size: 12px; letter-spacing: .16em; text-transform: uppercase; color: var(--accent); margin: 0 0 20px; font-weight: 500; }
  h1 { font-family: var(--font-display); font-weight: 600; font-size: clamp(44px, 7.5vw, 74px); line-height: 1.02; letter-spacing: -0.02em; margin: 0; text-wrap: balance; }
  .subline { margin: 18px 0 0; font-size: 19px; color: var(--ink-2); }
  .subline b { font-weight: 600; color: var(--ink); }
  .meta-row { margin-top: 8px; font-family: var(--font-mono); font-size: 13px; color: var(--muted); letter-spacing: .02em; }
  .actions { display: flex; flex-wrap: wrap; align-items: center; gap: 14px 22px; margin-top: 30px; }
  .btn { display: inline-flex; align-items: center; gap: 9px; background: var(--accent); color: var(--accent-ink); font-weight: 600; font-size: 15px; padding: 12px 22px; border-radius: 8px; box-shadow: var(--shadow); transition: transform .16s ease, filter .16s ease; }
  .btn:hover { transform: translateY(-1px); filter: brightness(1.05); }
  .btn.alt { background: transparent; color: var(--accent); border: 1px solid var(--line-strong); box-shadow: none; }
  .btn.alt:hover { background: var(--accent-soft); filter: none; }
  .links { display: flex; flex-wrap: wrap; gap: 20px; font-family: var(--font-mono); font-size: 13.5px; }
  .links a { color: var(--muted); border-bottom: 1px solid transparent; padding-bottom: 1px; transition: color .15s ease, border-color .15s ease; }
  .links a:hover { color: var(--accent); border-color: var(--accent); }
  section { padding: 46px 0; border-top: 1px solid var(--line); }
  section:first-of-type { border-top: 0; }
  .about p { font-family: var(--font-display); font-size: clamp(19px, 2.4vw, 23px); line-height: 1.55; color: var(--ink-2); margin: 0 0 18px; max-width: 62ch; text-wrap: pretty; }
  .about p:last-child { margin-bottom: 0; }
  .entry { display: grid; grid-template-columns: 140px 1fr; gap: 6px 30px; padding: 6px 0 30px; }
  .entry:last-child { padding-bottom: 0; }
  .when { font-family: var(--font-mono); font-size: 12.5px; color: var(--muted); font-variant-numeric: tabular-nums; letter-spacing: .01em; padding-top: 5px; line-height: 1.5; }
  .when .now { color: var(--accent); }
  .place { display: block; color: var(--faint); margin-top: 3px; }
  .role { font-family: var(--font-display); font-size: 21px; font-weight: 600; letter-spacing: -0.01em; margin: 0; }
  .org { color: var(--muted); font-size: 14.5px; margin: 3px 0 12px; }
  .org b { color: var(--ink-2); font-weight: 600; }
  .summary { margin: 0 0 12px; color: var(--ink-2); }
  ul.bul { list-style: none; margin: 0; padding: 0; display: flex; flex-direction: column; gap: 7px; }
  ul.bul li { position: relative; padding-left: 20px; color: var(--ink-2); }
  ul.bul li::before { content: ""; position: absolute; left: 3px; top: 11px; width: 5px; height: 5px; border: 1.5px solid var(--accent); border-radius: 50%; }
  .stack { margin-top: 14px; font-family: var(--font-mono); font-size: 12px; line-height: 1.9; color: var(--muted); letter-spacing: .01em; }
  .stack .k { color: var(--faint); text-transform: uppercase; letter-spacing: .14em; margin-right: 8px; }
  .rows { display: flex; flex-direction: column; gap: 20px; }
  .row { display: grid; grid-template-columns: 140px 1fr; gap: 8px 30px; }
  .row .k { font-family: var(--font-mono); font-size: 12.5px; color: var(--muted); text-transform: uppercase; letter-spacing: .08em; padding-top: 2px; }
  .row .v { color: var(--ink-2); }
  .row .v .deg { font-family: var(--font-display); font-size: 18px; color: var(--ink); font-weight: 600; }
  .row .v .sub { color: var(--muted); font-size: 14px; }
  footer { padding: 40px 0 76px; border-top: 1px solid var(--line); color: var(--faint); font-family: var(--font-mono); font-size: 12.5px; line-height: 1.8; }
  a:focus-visible, .btn:focus-visible, .theme-toggle:focus-visible { outline: 2px solid var(--accent); outline-offset: 3px; border-radius: 6px; }
  .theme-toggle { position: fixed; top: 16px; right: 16px; z-index: 20; width: 40px; height: 40px; border-radius: 50%; display: inline-flex; align-items: center; justify-content: center; background: var(--surface); color: var(--ink-2); border: 1px solid var(--line); box-shadow: var(--shadow); cursor: pointer; transition: color .16s ease, border-color .16s ease, transform .16s ease; }
  .theme-toggle:hover { color: var(--accent); border-color: var(--line-strong); transform: translateY(-1px); }
  .theme-toggle svg { display: block; }
  @keyframes rise { from { opacity: 0; transform: translateY(12px); } to { opacity: 1; transform: none; } }
  .anim { opacity: 0; animation: rise .6s cubic-bezier(.2,.6,.2,1) forwards; }
  .d1 { animation-delay: .05s; } .d2 { animation-delay: .13s; } .d3 { animation-delay: .21s; } .d4 { animation-delay: .29s; }
  @media (prefers-reduced-motion: reduce) { .anim { opacity: 1; animation: none; } html { scroll-behavior: auto; } }
  @media (max-width: 620px) {
    .entry, .row { grid-template-columns: 1fr; gap: 4px; }
    .when { padding-top: 0; }
    header.hero { padding-top: 56px; }
    .hero-grid { grid-template-columns: 1fr; gap: 26px; justify-items: start; }
    .portrait { width: 132px; height: 165px; order: -1; }
  }
"""

SITE_SCRIPT = """
  (function () {
    var root = document.documentElement;
    var btn = document.getElementById('theme-toggle');
    var MOON = '<svg viewBox="0 0 24 24" width="18" height="18" fill="none" stroke="currentColor" stroke-width="1.7" stroke-linecap="round" stroke-linejoin="round"><path d="M21 12.8A9 9 0 1 1 11.2 3a7 7 0 0 0 9.8 9.8Z"/></svg>';
    var SUN = '<svg viewBox="0 0 24 24" width="18" height="18" fill="none" stroke="currentColor" stroke-width="1.7" stroke-linecap="round" stroke-linejoin="round"><circle cx="12" cy="12" r="4"/><path d="M12 2v2M12 20v2M4.9 4.9l1.4 1.4M17.7 17.7l1.4 1.4M2 12h2M20 12h2M4.9 19.1l1.4-1.4M17.7 6.3l1.4-1.4"/></svg>';
    function isDark() { var t = root.getAttribute('data-theme'); if (t === 'dark') return true; if (t === 'light') return false; return window.matchMedia('(prefers-color-scheme: dark)').matches; }
    function paint() { btn.innerHTML = isDark() ? SUN : MOON; btn.setAttribute('aria-label', isDark() ? 'Switch to light theme' : 'Switch to dark theme'); }
    var saved = null; try { saved = localStorage.getItem('theme'); } catch (e) {}
    if (saved === 'dark' || saved === 'light') root.setAttribute('data-theme', saved);
    paint();
    btn.addEventListener('click', function () { var next = isDark() ? 'light' : 'dark'; root.setAttribute('data-theme', next); try { localStorage.setItem('theme', next); } catch (e) {} paint(); });
    window.matchMedia('(prefers-color-scheme: dark)').addEventListener('change', function () { if (!root.getAttribute('data-theme')) paint(); });
    if (window.matchMedia('(prefers-reduced-motion: reduce)').matches) { document.querySelectorAll('.anim').forEach(function (el) { el.style.opacity = 1; }); }
  })();
"""


def render_html(data, contact):
    name, title = data["name"], data["title"]
    exp0 = data["experience"][0]
    links = ['<a href="' + h(data["github"]) + '" target="_blank" rel="noopener">GitHub</a>']
    if contact["homepage"]:
        links.append('<a href="' + h(homepage_url(contact)) + '" target="_blank" rel="noopener">LinkedIn</a>')
    if contact["email"]:
        links.append('<a href="mailto:' + h(contact["email"]) + '">Email</a>')

    p = ['<!DOCTYPE html>\n<html lang="en">\n<head>\n'
         '<meta charset="utf-8" />\n<meta name="viewport" content="width=device-width, initial-scale=1" />\n'
         '<title>' + h(name) + ', ' + h(title) + '</title>\n'
         '<meta name="description" content="' + h(name) + ', ' + h(title) + ', based in ' + h(data["location"]) + '. Résumé and portfolio." />\n'
         '<style>' + SITE_CSS + '</style>\n' + render_jsonld(data, contact) + '</head>\n<body>\n']
    p.append('<button class="theme-toggle" id="theme-toggle" type="button" aria-label="Toggle colour theme"></button>\n\n')
    p.append('  <div class="wrap">\n')
    p.append(
        '    <header class="hero">\n      <div class="hero-grid">\n        <div class="hero-text">\n'
        '          <p class="kicker anim d1">' + h(data["eyebrow"]) + '</p>\n'
        '          <h1 class="anim d1">' + h(name) + '</h1>\n'
        '          <p class="subline anim d2"><b>' + h(exp0["role"]) + '</b> at ' + h(exp0["org"]) + ', ' + h(data["tagline"]) + '</p>\n'
        '          <p class="meta-row anim d2">' + h(data["location"]) + ' · ' + h(data["country"]) + '</p>\n'
        '          <div class="actions anim d3">\n'
        '            <a class="btn" href="resume.pdf" target="_blank" rel="noopener">Open PDF résumé ↗</a>\n'
        '            <a class="btn alt" href="resume.docx">Word (.docx) ↓</a>\n'
        '            <nav class="links">' + " ".join(links) + '</nav>\n'
        '          </div>\n        </div>\n'
        '        <img class="portrait anim d2" src="' + h(data["photo"]) + '" alt="' + h(name) + '" width="172" height="215" />\n'
        '      </div>\n    </header>\n\n')
    p.append('    <section class="about anim d4">\n')
    for para in data["about"]:
        p.append('      <p>' + h(para) + '</p>\n')
    p.append('    </section>\n\n')
    p.append('    <section>\n      <div class="rule-head"><span class="label">Experience</span></div>\n')
    for e in data["experience"]:
        place = h(e["city"]) + ", " + ccode(e["country"])
        if is_present(e["end"]):
            when = '<span class="now">' + h(e["start"]) + ' — Now</span><span class="place">' + place + '</span>'
        else:
            when = h(e["start"]) + ' — ' + h(e["end"]) + '<span class="place">' + place + '</span>'
        p.append('\n      <div class="entry">\n        <div class="when">' + when + '</div>\n        <div class="what">\n'
                 '          <h3 class="role">' + h(e["role"]) + '</h3>\n'
                 '          <p class="org"><b>' + h(e["org"]) + '</b></p>\n'
                 '          <p class="summary">' + h(e["summary"]) + '</p>\n')
        if e.get("bullets"):
            p.append('          <ul class="bul">\n')
            for b in e["bullets"]:
                p.append('            <li>' + h(b) + '</li>\n')
            p.append('          </ul>\n')
        if e.get("tech"):
            p.append('          <p class="stack"><span class="k">Stack</span>' + h(stack(e["tech"])) + '</p>\n')
        p.append('        </div>\n      </div>\n')
    p.append('    </section>\n\n')
    p.append('    <section>\n      <div class="rule-head"><span class="label">Skills</span></div>\n      <div class="rows">\n')
    for sk in data["skills"]:
        p.append('        <div class="row">\n          <div class="k">' + h(sk["label"]) + '</div>\n          <div class="v">' + h(sk["items"]) + '</div>\n        </div>\n')
    p.append('      </div>\n    </section>\n\n')
    p.append('    <section>\n      <div class="rule-head"><span class="label">Education</span></div>\n      <div class="rows">\n')
    for ed in data["education"]:
        p.append('        <div class="row">\n          <div class="k">' + h(ed["start"]) + ' — ' + h(ed["end"]) + '</div>\n'
                 '          <div class="v"><span class="deg">' + h(ed["degree"]) + '</span><br><span class="sub">' + h(ed["org"]) + ' · ' + h(ed["detail"]) + '</span></div>\n        </div>\n')
    p.append('      </div>\n    </section>\n\n')
    p.append('    <footer>\n      References available on request.\n    </footer>\n')
    p.append('  </div>\n\n<script>' + SITE_SCRIPT + '</script>\n</body>\n</html>\n')
    return "".join(p)


# ===========================================================================
# Print document (PDF via WeasyPrint)
# ===========================================================================

PRINT_CSS = """
  @page { size: A4; margin: 15mm 15mm 13mm; }
  * { box-sizing: border-box; }
  body { margin: 0; font-family: "IBM Plex Sans", ui-sans-serif, system-ui, sans-serif; color: #15181f; font-size: 10.2pt; line-height: 1.42; }
  a { color: #6a7180; text-decoration: underline; text-decoration-color: #bfc4cc; }
  h1 { font-family: "IBM Plex Serif", ui-serif, Georgia, serif; font-weight: 600; font-size: 25pt; line-height: 1.05; letter-spacing: -0.01em; margin: 0; }
  .role { font-family: "IBM Plex Mono", ui-monospace, monospace; font-size: 8pt; letter-spacing: .16em; text-transform: uppercase; color: #b0651a; margin: 6pt 0 0; font-weight: 500; }
  .contact { font-family: "IBM Plex Mono", ui-monospace, monospace; font-size: 8pt; line-height: 1.6; color: #6a7180; margin: 7pt 0 0; }
  .hr { height: 1.6pt; background: #b0651a; margin: 9pt 0 0; border: 0; }
  .summary { font-family: "IBM Plex Serif", ui-serif, Georgia, serif; font-size: 10.5pt; line-height: 1.45; color: #3a4049; margin: 11pt 0 0; }
  h2 { font-family: "IBM Plex Mono", ui-monospace, monospace; font-size: 8pt; letter-spacing: .2em; text-transform: uppercase; color: #6a7180; font-weight: 600; border-bottom: 0.6pt solid #d3d8df; padding-bottom: 5pt; margin: 15pt 0 9pt; break-after: avoid; }
  .entry { margin-bottom: 10pt; break-inside: avoid; }
  h3 { font-family: "IBM Plex Serif", ui-serif, Georgia, serif; font-size: 11.5pt; font-weight: 600; margin: 0; letter-spacing: -0.01em; }
  h3 .org { font-family: "IBM Plex Sans", sans-serif; font-weight: 600; font-size: 9.5pt; color: #6a7180; }
  .meta { font-family: "IBM Plex Mono", ui-monospace, monospace; font-size: 7.8pt; color: #6a7180; margin: 2pt 0 0; }
  .meta .now { color: #b0651a; }
  .sum { margin: 4pt 0 5pt; font-size: 9.4pt; line-height: 1.4; color: #3a4049; }
  ul { list-style: none; margin: 0; padding: 0; }
  li { padding-left: 11pt; text-indent: -11pt; font-size: 9.4pt; line-height: 1.38; color: #3a4049; margin-bottom: 2.5pt; }
  li::before { content: ""; display: inline-block; width: 3pt; height: 3pt; border: 1pt solid #b0651a; border-radius: 50%; margin: 0 6pt 1pt 2pt; vertical-align: middle; }
  .stack { margin-top: 6pt; font-family: "IBM Plex Mono", ui-monospace, monospace; font-size: 7.6pt; color: #6a7180; line-height: 1.6; }
  .stack b { color: #949aa6; font-weight: 600; letter-spacing: .1em; text-transform: uppercase; margin-right: 6pt; }
  .skill { font-size: 9.4pt; color: #3a4049; margin-bottom: 3pt; }
  .skill b { font-family: "IBM Plex Mono", ui-monospace, monospace; font-size: 7.8pt; text-transform: uppercase; letter-spacing: .08em; color: #6a7180; font-weight: 600; margin-right: 7pt; }
  .edu { margin-bottom: 7pt; break-inside: avoid; }
  .edu .deg { font-family: "IBM Plex Serif", ui-serif, Georgia, serif; font-size: 10.5pt; font-weight: 600; color: #15181f; }
  .edu .meta2 { font-family: "IBM Plex Mono", ui-monospace, monospace; font-size: 7.8pt; color: #6a7180; margin: 1pt 0 0; }
"""


def render_print_html(data, contact, json_href="resume.json"):
    name = data["name"]
    parts = []
    if contact["phone"]:
        parts.append(h(contact["phone"]))
    if contact["email"]:
        parts.append('<a href="mailto:' + h(contact["email"]) + '">' + h(contact["email"]) + '</a>')
    if contact["homepage"]:
        hp = strip_scheme(contact["homepage"])
        parts.append('<a href="' + h(homepage_url(contact)) + '">' + h(hp) + '</a>')
    parts.append('<a href="' + h(data["github"]) + '">' + h(strip_scheme(data["github"])) + '</a>')
    contact_line = h(data["location"]) + ", " + h(data["country"]) + "  ·  " + "  ·  ".join(parts)

    summary = " ".join(data["about"])
    keywords = h(", ".join(skill_keywords(data)))
    exp0 = data["experience"][0]
    desc = h(name) + ", " + h(data["title"]) + ", " + h(exp0["role"]) + " at " + h(exp0["org"]) + "."

    p = ['<!DOCTYPE html>\n<html lang="en">\n<head>\n<meta charset="utf-8" />\n'
         '<title>' + h(name) + ' Résumé</title>\n'
         '<meta name="author" content="' + h(name) + '" />\n'
         '<meta name="description" content="' + desc + '" />\n'
         '<meta name="keywords" content="' + keywords + '" />\n'
         '<meta name="generator" content="build.py" />\n'
         '<link rel="attachment" href="' + h(json_href) + '" title="resume.json" />\n'
         '<style>' + PRINT_CSS + '</style>\n</head>\n<body>\n']
    p.append('  <header>\n    <h1>' + h(name) + '</h1>\n'
             '    <p class="role">' + h(data["eyebrow"]) + '</p>\n'
             '    <p class="contact">' + contact_line + '</p>\n  </header>\n'
             '  <hr class="hr" />\n')
    p.append('  <p class="summary">' + h(summary) + '</p>\n')

    p.append('  <section>\n    <h2>Experience</h2>\n')
    for e in data["experience"]:
        loc = h(e["city"]) + ", " + ccode(e["country"])
        if is_present(e["end"]):
            dates = '<span class="now">' + h(e["start"]) + ' — Now</span>'
        else:
            dates = h(e["start"]) + ' — ' + h(e["end"])
        p.append('    <article class="entry">\n'
                 '      <h3>' + h(e["role"]) + '<span class="org">, ' + h(e["org"]) + '</span></h3>\n'
                 '      <p class="meta">' + dates + '  ·  ' + loc + '</p>\n'
                 '      <p class="sum">' + h(e["summary"]) + '</p>\n')
        if e.get("bullets"):
            p.append('      <ul>\n')
            for b in e["bullets"]:
                p.append('        <li>' + h(b) + '</li>\n')
            p.append('      </ul>\n')
        if e.get("tech"):
            p.append('      <p class="stack"><b>Stack</b>' + h(stack(e["tech"])) + '</p>\n')
        p.append('    </article>\n')
    p.append('  </section>\n')

    p.append('  <section>\n    <h2>Skills</h2>\n')
    for sk in data["skills"]:
        p.append('    <p class="skill"><b>' + h(sk["label"]) + '</b>' + h(sk["items"]) + '</p>\n')
    p.append('  </section>\n')

    p.append('  <section>\n    <h2>Education</h2>\n')
    for ed in data["education"]:
        p.append('    <div class="edu">\n'
                 '      <span class="deg">' + h(ed["degree"]) + '</span>, ' + h(ed["org"]) + '\n'
                 '      <p class="meta2">' + h(ed["start"]) + ' — ' + h(ed["end"]) + '  ·  ' + h(ed["detail"]) + '</p>\n'
                 '    </div>\n')
    p.append('  </section>\n')

    p.append('</body>\n</html>\n')
    return "".join(p)


# ===========================================================================
# Word document (.docx)
# ===========================================================================

def build_docx(data, contact):
    """Build a Word résumé as a python-docx Document, single column and in order."""
    from docx import Document
    from docx.shared import Pt, RGBColor

    accent = RGBColor(0xB0, 0x65, 0x1E)
    muted = RGBColor(0x6A, 0x71, 0x80)
    doc = Document()
    normal = doc.styles["Normal"].font
    normal.name = "Calibri"
    normal.size = Pt(10.5)

    name = doc.add_paragraph().add_run(data["name"])
    name.bold = True
    name.font.size = Pt(22)

    role = doc.add_paragraph().add_run(data["eyebrow"])
    role.bold = True
    role.font.size = Pt(9)
    role.font.color.rgb = accent

    bits = [data["location"] + ", " + data["country"]]
    if contact["phone"]:
        bits.append(contact["phone"])
    if contact["email"]:
        bits.append(contact["email"])
    if contact["homepage"]:
        bits.append(strip_scheme(contact["homepage"]))
    bits.append(strip_scheme(data["github"]))
    c = doc.add_paragraph().add_run(" | ".join(bits))
    c.font.size = Pt(9)
    c.font.color.rgb = muted

    doc.add_paragraph(" ".join(data["about"]))

    def dated(text):
        run = doc.add_paragraph().add_run(text)
        run.italic = True
        run.font.size = Pt(9)
        run.font.color.rgb = muted

    doc.add_heading("Experience", level=1)
    for e in data["experience"]:
        end = "Present" if is_present(e["end"]) else e["end"]
        doc.add_paragraph().add_run(e["role"] + ", " + e["org"]).bold = True
        dated(str(e["start"]) + " - " + str(end) + " | " + e["city"] + ", " + e["country"])
        doc.add_paragraph(e["summary"])
        for b in (e.get("bullets") or []):
            doc.add_paragraph(b, style="List Bullet")
        if e.get("tech"):
            tech = doc.add_paragraph().add_run("Tech: " + ", ".join(split_commas(e["tech"])))
            tech.font.size = Pt(9)
            tech.font.color.rgb = muted

    doc.add_heading("Skills", level=1)
    for sk in data["skills"]:
        p = doc.add_paragraph()
        p.add_run(sk["label"] + ": ").bold = True
        p.add_run(sk["items"])

    doc.add_heading("Education", level=1)
    for ed in data["education"]:
        doc.add_paragraph().add_run(ed["degree"] + ", " + ed["org"]).bold = True
        dated(str(ed["start"]) + " - " + str(ed["end"]) + " | " + ed["detail"])

    return doc


def main():
    data = yaml.safe_load((HERE / "content.yaml").read_text())
    contact = read_contact()
    (HERE / "index.html").write_text(render_html(data, contact))
    # resume.json is written first so the PDF can attach it.
    (HERE / "resume.json").write_text(render_json_resume(data, contact))
    (HERE / "resume.txt").write_text(render_text(data, contact))
    (HERE / "llms.txt").write_text(render_llms(data, contact))
    print_html = render_print_html(data, contact)
    (HERE / "resume-print.html").write_text(print_html)
    made = ["index.html", "resume.json", "resume.txt", "llms.txt", "resume-print.html"]
    try:
        build_docx(data, contact).save(str(HERE / "resume.docx"))
        made.append("resume.docx")
    except ImportError:
        pass
    try:
        from weasyprint import HTML
        HTML(string=print_html, base_url=str(HERE)).write_pdf(
            str(HERE / "resume.pdf"), pdf_variant="pdf/ua-1")
        made.append("resume.pdf")
    except ImportError:
        pass
    print("Generated " + ", ".join(made))


if __name__ == "__main__":
    main()
